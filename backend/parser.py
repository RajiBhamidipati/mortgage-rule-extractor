"""Document parsing for PDF and Word files.

Extracts text with section/page tracking, table content,
and overlapping chunk boundaries per PRD F-01 through F-04.
"""

import re
import pdfplumber
from docx import Document
from dataclasses import dataclass, field


@dataclass
class ParsedSection:
    """A section of a parsed document."""
    section_heading: str | None
    text: str
    page: int | None  # None for Word docs
    char_count: int
    tables: list[str] = field(default_factory=list)  # Table text representations


@dataclass
class ParsedDocument:
    """Complete parsed document with sections and metadata."""
    doc_name: str
    doc_type: str  # "pdf" or "docx"
    total_pages: int | None
    sections: list[ParsedSection]
    full_text: str  # Concatenated text for Claude prompt
    definitions: list[str] = field(default_factory=list)  # Global context summary (F-03)


def parse_pdf(file_path: str, doc_name: str) -> ParsedDocument:
    """Extract text from PDF with page-level tracking and table extraction."""
    sections: list[ParsedSection] = []

    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            # Extract main text
            text = page.extract_text() or ""

            # Extract tables as formatted text
            table_texts = []
            for table in page.extract_tables():
                if table:
                    rows = []
                    for row in table:
                        cells = [str(cell or "").strip() for cell in row]
                        rows.append(" | ".join(cells))
                    table_texts.append("\n".join(rows))

            # Detect section heading (first short uppercase-ish line or numbered heading)
            heading = _detect_heading(text)

            # Combine text and tables
            combined = text
            if table_texts:
                combined += "\n\n[TABLE]\n" + "\n[/TABLE]\n\n[TABLE]\n".join(table_texts) + "\n[/TABLE]"

            sections.append(ParsedSection(
                section_heading=heading,
                text=combined,
                page=i + 1,
                char_count=len(combined),
                tables=table_texts,
            ))

    full_text = "\n\n".join(s.text for s in sections)
    definitions = _extract_definitions(full_text)

    return ParsedDocument(
        doc_name=doc_name,
        doc_type="pdf",
        total_pages=len(sections),
        sections=sections,
        full_text=full_text,
        definitions=definitions,
    )


def parse_docx(file_path: str, doc_name: str) -> ParsedDocument:
    """Extract text from Word doc with heading-based section tracking."""
    doc = Document(file_path)
    sections: list[ParsedSection] = []
    current_heading: str | None = None
    current_text_parts: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]  # Strip namespace

        if tag == "p":
            # It's a paragraph
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            if para.style and para.style.name and para.style.name.startswith("Heading"):
                # Flush previous section
                if current_text_parts:
                    text = "\n".join(current_text_parts)
                    sections.append(ParsedSection(
                        section_heading=current_heading,
                        text=text,
                        page=None,
                        char_count=len(text),
                    ))
                    current_text_parts = []
                current_heading = para.text
            else:
                if para.text.strip():
                    current_text_parts.append(para.text)

        elif tag == "tbl":
            # It's a table — extract cell contents
            for tbl in doc.tables:
                if tbl._element is element:
                    table_rows = []
                    for row in tbl.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        table_rows.append(" | ".join(cells))
                    table_text = "\n".join(table_rows)
                    current_text_parts.append(f"\n[TABLE]\n{table_text}\n[/TABLE]")
                    break

    # Flush final section
    if current_text_parts:
        text = "\n".join(current_text_parts)
        sections.append(ParsedSection(
            section_heading=current_heading,
            text=text,
            page=None,
            char_count=len(text),
        ))

    full_text = "\n\n".join(s.text for s in sections)
    definitions = _extract_definitions(full_text)

    return ParsedDocument(
        doc_name=doc_name,
        doc_type="docx",
        total_pages=None,
        sections=sections,
        full_text=full_text,
        definitions=definitions,
    )


def parse_document(file_path: str, doc_name: str) -> ParsedDocument:
    """Parse a document based on file extension."""
    lower = file_path.lower()
    if lower.endswith(".pdf"):
        return parse_pdf(file_path, doc_name)
    elif lower.endswith(".docx"):
        return parse_docx(file_path, doc_name)
    else:
        raise ValueError(f"Unsupported file type: {file_path}. Supported: .pdf, .docx")


def _detect_heading(page_text: str) -> str | None:
    """Heuristic heading detection from page text."""
    lines = page_text.strip().split("\n")
    if not lines:
        return None

    first_line = lines[0].strip()

    # Numbered heading pattern: "1. Something" or "1.2 Something"
    if re.match(r"^\d+(\.\d+)*\.?\s+\S", first_line) and len(first_line) < 100:
        return first_line

    # ALL CAPS heading
    if len(first_line) < 80 and first_line == first_line.upper() and len(first_line) > 3:
        return first_line

    return None


def _extract_definitions(full_text: str) -> list[str]:
    """Extract definitions, acronyms, and glossary entries for Global Context Summary (F-03).

    Looks for patterns like:
    - "X — definition"
    - "X: definition"
    - Lines in an Appendix/Definitions section
    """
    definitions = []

    # Find lines with em-dash or colon-based definitions
    for line in full_text.split("\n"):
        line = line.strip()
        # Pattern: "ACRONYM — definition" or "ACRONYM: definition"
        if re.match(r"^[A-Z]{2,}[\s]*[—–\-:]\s+.{10,}", line):
            definitions.append(line)

    return definitions


def build_chunk_with_context(section: ParsedSection, definitions: list[str], doc_name: str) -> str:
    """Build a text chunk for Claude extraction with Global Context Summary prepended (F-03)."""
    parts = []

    # Global Context Summary
    if definitions:
        parts.append("=== GLOBAL CONTEXT: DEFINITIONS AND ACRONYMS ===")
        for d in definitions:
            parts.append(f"  {d}")
        parts.append("=== END GLOBAL CONTEXT ===\n")

    # Section content
    if section.section_heading:
        parts.append(f"SECTION: {section.section_heading}")
    parts.append(section.text)

    return "\n".join(parts)
