"""Generate a synthetic lender policy document for development and testing.
Run: python create_sample_policy.py
Outputs: sample_lender_policy.docx and sample_lender_policy.pdf (if docx2pdf available)
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading("Meridian Building Society", level=0)
subtitle = doc.add_heading("Mortgage Lending Criteria — Version 2.1", level=1)
doc.add_paragraph("Effective Date: 1 January 2026")
doc.add_paragraph("Classification: Internal — Confidential")
doc.add_paragraph("This document sets out the lending criteria for all residential and buy-to-let mortgage products offered by Meridian Building Society. All applications must satisfy the criteria below unless a specific exception is approved by the Senior Underwriting Panel.")
doc.add_paragraph("")

# ── Section 1: LTV ──
doc.add_heading("1. Loan-to-Value (LTV) Requirements", level=1)
doc.add_paragraph(
    "The maximum loan-to-value ratio for standard residential mortgages is 95%. "
    "For buy-to-let mortgages, the maximum LTV is 75%. "
    "Applications where the LTV exceeds 90% require private mortgage insurance (PMI) arranged through an approved provider."
)
doc.add_paragraph(
    "New build properties are subject to a reduced maximum LTV of 85%. "
    "For houses in multiple occupation (HMOs), the maximum LTV is 70%."
)
doc.add_paragraph(
    "The maximum loan amount for any single application is £2,000,000. "
    "Loans above £1,000,000 are classified as high-value and require Senior Underwriter approval."
)
# Table: LTV summary
table = doc.add_table(rows=6, cols=3)
table.style = "Light Grid Accent 1"
headers = ["Property Type", "Max LTV", "Notes"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ("Standard Residential", "95%", "PMI required above 90%*"),
    ("Buy-to-Let", "75%", "Minimum rental coverage 125%"),
    ("New Build", "85%", "Developer must be on approved list"),
    ("HMO", "70%", "Maximum 6 units"),
    ("Interest Only", "75%", "Repayment vehicle required"),
]
for row_idx, (ptype, ltv, notes) in enumerate(data, 1):
    table.rows[row_idx].cells[0].text = ptype
    table.rows[row_idx].cells[1].text = ltv
    table.rows[row_idx].cells[2].text = notes

doc.add_paragraph("")
doc.add_paragraph("* PMI premium is added to the monthly payment and is non-refundable for the first 24 months.", style="Intense Quote")

# ── Section 2: Income ──
doc.add_heading("2. Income Requirements", level=1)
doc.add_paragraph(
    "The minimum gross annual income for a single applicant is £25,000. "
    "For joint applications, the combined minimum gross annual income is £30,000."
)
doc.add_paragraph(
    "The maximum income multiple is 4.5x gross annual income for single applicants. "
    "For joint applicants, the maximum income multiple is 3.75x combined gross annual income. "
    "First-time buyers with a deposit of 15% or more may be offered an enhanced income multiple of up to 4.75x, subject to affordability assessment."
)
doc.add_paragraph(
    "Acceptable income types include: basic salary, guaranteed overtime, regular commission (minimum 12 months history), "
    "and pension income. Discretionary bonuses may be considered at 50% of the two-year average. "
    "Rental income from existing buy-to-let properties is assessed at 75% of the gross rental figure."
)

# ── Section 3: Employment ──
doc.add_heading("3. Employment Criteria", level=1)
doc.add_paragraph(
    "Applicants in permanent employment must have a minimum of 6 months continuous service with their current employer. "
    "Applicants still within a probationary period are not accepted."
)
doc.add_paragraph(
    "Self-employed applicants, including sole traders and limited company directors, must provide a minimum of 2 years "
    "of trading history evidenced by SA302 tax calculations or certified accounts. Income is assessed as the lower of "
    "the two-year average or the most recent year."
)
doc.add_paragraph(
    "Contract workers must demonstrate a minimum of 12 months continuous contracting history in the same sector. "
    "Day-rate contractors are assessed on an annualised basis: day rate × 5 × 46 weeks. "
    "Zero-hours contract workers are not accepted."
)
doc.add_paragraph(
    "Retired applicants receiving a guaranteed pension income are accepted provided the pension income meets the minimum income threshold "
    "and the mortgage term does not extend beyond the applicant's 80th birthday."
)

# ── Section 4: Credit ──
doc.add_heading("4. Credit History Criteria", level=1)
doc.add_paragraph(
    "The minimum acceptable credit score is 620 (Experian). Applicants with a credit score below 620 will be declined."
)
doc.add_paragraph(
    "No County Court Judgements (CCJs) are accepted within the last 6 years, regardless of whether they have been satisfied. "
    "No bankruptcy, Individual Voluntary Arrangements (IVAs), or Debt Relief Orders within the last 6 years."
)
doc.add_paragraph(
    "A maximum of 1 missed mortgage or secured loan payment is permitted in the last 12 months. "
    "A maximum of 2 missed unsecured credit payments are permitted in the last 12 months. "
    "Applicants with 3 or more missed payments of any type in the last 12 months will be declined."
)
doc.add_paragraph(
    "Total unsecured debt must not exceed 15% of gross annual income at the time of application."
)

# ── Section 5: Property ──
doc.add_heading("5. Property Requirements", level=1)
doc.add_paragraph(
    "The minimum property value is £75,000. Properties valued below this threshold are outside lending criteria."
)
doc.add_paragraph(
    "Properties of non-standard construction (e.g. steel frame, concrete panel, timber frame without brick skin) "
    "are considered on a case-by-case basis and must be referred to the specialist underwriting team."
)
doc.add_paragraph(
    "Leasehold properties must have a minimum unexpired lease term of 70 years at the time of application, "
    "or 40 years beyond the end of the mortgage term, whichever is greater."
)
doc.add_paragraph(
    "The minimum acceptable EPC rating is E for standard residential properties. "
    "Properties with an EPC rating of F or G are not accepted unless the applicant commits to remediation works within 12 months "
    "and an appropriate retention is applied to the advance."
)
doc.add_paragraph(
    "Flats above commercial premises are accepted where the commercial element does not exceed 40% of the total floor area. "
    "Properties above fast food outlets, nightclubs, or petrol stations are not accepted."
)

# ── Section 6: Affordability ──
doc.add_heading("6. Affordability Assessment", level=1)
doc.add_paragraph(
    "All applications are subject to an affordability stress test at the higher of: "
    "(a) the lender's standard variable rate (SVR) plus 3 percentage points, or "
    "(b) a floor rate of 7%. "
    "The applicant must demonstrate sufficient disposable income to meet the stressed payment."
)
doc.add_paragraph(
    "The maximum debt-to-income (DTI) ratio is 45%. This includes all committed monthly expenditure: "
    "existing mortgage payments, personal loans, credit card minimum payments, hire purchase agreements, "
    "child maintenance, and student loan repayments."
)
doc.add_paragraph(
    "Childcare costs must be verified and included in the affordability calculation. "
    "The Office for National Statistics (ONS) expenditure data is used as a benchmark for essential living costs."
)

# ── Section 7: Applicant ──
doc.add_heading("7. Applicant Eligibility", level=1)
doc.add_paragraph(
    "The minimum age at the time of application is 18. "
    "The maximum age at the end of the mortgage term is 70 for standard products and 80 for retirement interest-only products."
)
doc.add_paragraph(
    "Applicants must be UK residents or hold Indefinite Leave to Remain (settled status). "
    "EEA nationals with pre-settled status are accepted for terms up to 5 years or until settled status is confirmed. "
    "Applicants on a Tier 2 (Skilled Worker) visa are considered where the visa has at least 3 years remaining."
)
doc.add_paragraph(
    "A maximum of 2 applicants are permitted per application. "
    "All applicants must be named on the property title."
)

# ── Section 8: Loan Terms ──
doc.add_heading("8. Loan Structure and Terms", level=1)
doc.add_paragraph(
    "The minimum mortgage term is 5 years. The maximum mortgage term is 35 years. "
    "Interest-only mortgages are available up to a maximum LTV of 50%, with an approved repayment vehicle required "
    "(e.g. ISA, pension, investment portfolio, or sale of another property)."
)
doc.add_paragraph(
    "Part-and-part mortgages (a combination of repayment and interest-only) are permitted. "
    "The interest-only element must not exceed 50% of the total loan amount."
)
doc.add_paragraph(
    "Early repayment charges (ERCs) apply during the initial product period. "
    "ERC rates: Year 1 — 5%, Year 2 — 4%, Year 3 — 3%, Year 4 — 2%, Year 5 — 1%. "
    "Overpayments of up to 10% of the outstanding balance per annum are permitted without charge."
)
doc.add_paragraph(
    "Capital raising is permitted up to a maximum of £500,000 per application, "
    "subject to a maximum LTV of 80% and a satisfactory explanation of purpose. "
    "Capital raising for business purposes is not permitted on residential mortgages."
)

# ── Appendix: Definitions ──
doc.add_heading("Appendix A: Definitions", level=1)
doc.add_paragraph("LTV — Loan-to-Value ratio: the loan amount expressed as a percentage of the property value.")
doc.add_paragraph("PMI — Private Mortgage Insurance: insurance that protects the lender against loss in the event of borrower default on high-LTV loans.")
doc.add_paragraph("HMO — House in Multiple Occupation: a property rented to 3 or more tenants forming 2 or more households.")
doc.add_paragraph("DTI — Debt-to-Income ratio: total monthly debt payments expressed as a percentage of gross monthly income.")
doc.add_paragraph("SVR — Standard Variable Rate: the lender's default interest rate applied after any initial product period ends.")
doc.add_paragraph("EPC — Energy Performance Certificate: a rating of the energy efficiency of a property from A (most efficient) to G (least efficient).")
doc.add_paragraph("SA302 — HMRC tax calculation: an official document showing an individual's taxable income for a given tax year.")
doc.add_paragraph("ERC — Early Repayment Charge: a fee charged when a borrower repays all or part of the mortgage during the initial product period.")
doc.add_paragraph("IVA — Individual Voluntary Arrangement: a formal agreement with creditors to repay debts over a fixed period.")

# Save
output_path = "/Users/Raji.Bhamidipati@finova.tech/mortgage-rule-extractor/sample_docs/sample_lender_policy.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
